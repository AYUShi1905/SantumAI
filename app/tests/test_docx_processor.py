import os
import sys
import io
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL

# Add app directory to path
sys.path.append(os.path.join(os.getcwd(), "app"))

from services.processor import DocumentProcessorService

def create_test_docx():
    doc = Document()
    doc.add_paragraph("Paragraph 1 before table.")
    
    # Add a 3x3 table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Fill header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Header 1'
    hdr_cells[1].text = 'Header 2'
    hdr_cells[2].text = 'Header 3'
    
    # Merged cell in row 1 (columns 0 and 1)
    row1_cells = table.rows[1].cells
    a = row1_cells[0]
    b = row1_cells[1]
    a.merge(b)
    a.text = "Merged Content"
    row1_cells[2].text = "Row 1 Col 2"
    
    # Standard row 2
    row2_cells = table.rows[2].cells
    row2_cells[0].text = "R2 C0"
    row2_cells[1].text = "R2 C1"
    row2_cells[2].text = "R2 C2"
    
    doc.add_paragraph("Paragraph 2 after table.")
    
    # Save to bytes
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()

def test_processor():
    processor = DocumentProcessorService()
    docx_content = create_test_docx()
    
    extracted_text = processor.process_docx(docx_content)
    
    print("--- Extracted Text ---")
    print(extracted_text)
    print("----------------------")
    
    # Assertions
    assert "Paragraph 1 before table." in extracted_text
    assert "Paragraph 2 after table." in extracted_text
    # Check for Markdown table structure
    assert "| Header 1 | Header 2 | Header 3 |" in extracted_text
    assert "| --- | --- | --- |" in extracted_text
    # Check for repeated text in merged cells
    assert "| Merged Content | Merged Content | Row 1 Col 2 |" in extracted_text
    assert "| R2 C0 | R2 C1 | R2 C2 |" in extracted_text
    
    print("Tests passed successfully!")

if __name__ == "__main__":
    test_processor()
